import streamlit as st
from docx import Document
from io import BytesIO

# CTR data
ctr_data = {
    1: 33.25, 2: 14.17, 3: 9.17, 4: 5.44, 5: 3.36,
    6: 2.18, 7: 1.49, 8: 1.11, 9: 0.83, 10: 0.65,
    11: 0.59, 12: 0.58, 13: 0.69, 14: 0.73, 15: 0.72,
    16: 0.60, 17: 0.70, 18: 0.49, 19: 0.54, 20: 0.46
}

def calculate_traffic_and_conversions(search_volume, current_rank, target_rank, app_start_rate, app_submit_rate):
    current_ctr = ctr_data.get(current_rank, 0) / 100
    target_ctr = ctr_data.get(target_rank, 0) / 100
    
    current_traffic = search_volume * current_ctr
    target_traffic = search_volume * target_ctr
    incremental_traffic = target_traffic - current_traffic
    
    current_app_starts = current_traffic * app_start_rate
    current_app_submits = current_app_starts * app_submit_rate
    
    target_app_starts = target_traffic * app_start_rate
    target_app_submits = target_app_starts * app_submit_rate
    
    incremental_app_starts = target_app_starts - current_app_starts
    incremental_app_submits = target_app_submits - current_app_submits
    
    return {
        "current_traffic": round(current_traffic),
        "target_traffic": round(target_traffic),
        "incremental_traffic": round(incremental_traffic),
        "current_app_starts": round(current_app_starts),
        "current_app_submits": round(current_app_submits),
        "target_app_starts": round(target_app_starts),
        "target_app_submits": round(target_app_submits),
        "incremental_app_starts": round(incremental_app_starts),
        "incremental_app_submits": round(incremental_app_submits)
    }

def main():
    st.set_page_config(page_title="SEO Content Optimizer", layout="wide")
    st.title("SEO Content Optimizer")
    st.write("Created by Brandon Lazovic")
    
    st.markdown("""
    ### How to use this tool:
    1. Fill in the information for each section below.
    2. Use the tooltips (?) for guidance on SEO best practices.
    3. Click the download button to generate a report in Word format.
    """)

    if 'data' not in st.session_state:
        st.session_state.data = {}

    keyword_research()
    serp_analysis()
    on_page_elements()
    internal_links()
    traffic_and_conversions()

    st.download_button(
        label="Download Word Report",
        data=get_word_download(),
        file_name="seo_content_optimization_report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

def keyword_research():
    st.header("Keyword Research")
    
    if 'keyword_data' not in st.session_state:
        st.session_state.keyword_data = [{"keyword": "", "search_volume": 0, "current_rank": 1}]
    
    def add_keyword():
        st.session_state.keyword_data.append({"keyword": "", "search_volume": 0, "current_rank": 1})
    
    def remove_keyword(index):
        st.session_state.keyword_data.pop(index)
    
    for i, keyword in enumerate(st.session_state.keyword_data):
        col1, col2, col3, col4 = st.columns([3, 2, 2, 1])
        
        with col1:
            keyword["keyword"] = st.text_input(f"Keyword {i+1}", keyword["keyword"], key=f"keyword_{i}")
        
        with col2:
            keyword["search_volume"] = st.number_input(f"Search Volume {i+1}", min_value=0, value=keyword["search_volume"], key=f"volume_{i}")
        
        with col3:
            keyword["current_rank"] = st.number_input(f"Current Rank {i+1}", min_value=1, max_value=100, value=keyword["current_rank"], key=f"rank_{i}")
        
        with col4:
            if st.button("Remove", key=f"remove_{i}"):
                remove_keyword(i)
                st.experimental_rerun()
    
    if st.button("Add Keyword"):
        add_keyword()
        st.experimental_rerun()

def serp_analysis():
    st.header("SERP Analysis")
    
    st.session_state.data['Search Intent For Target Keyword'] = st.text_area(
        "Search Intent For Target Keyword",
        help="Describe the search intent for the primary keyword."
    )
    
    st.session_state.data['Search Result Features For Target Keyword'] = st.text_area(
        "Search Result Features For Target Keyword",
        help="List the search features present in the SERP for the primary keyword."
    )
    
    st.session_state.data['FAQs Present Within Search Results'] = st.text_area(
        "FAQs Present Within Search Results",
        help="List common questions asked by users related to the primary keyword."
    )

def on_page_elements():
    st.header("On-Page Elements")
    
    st.session_state.data['Recommended Title (H1)'] = st.text_input(
        "Recommended Title (H1)",
        help="Enter the H1 tag. Ensure the primary keyword is targeted and it's 65 characters or less.",
        max_chars=65
    )
    
    st.session_state.data['Recommended Page Title'] = st.text_input(
        "Recommended Page Title",
        help="Enter the meta title. Ensure the primary keyword is targeted and it's 65 characters or less.",
        max_chars=65
    )
    
    st.session_state.data['Recommended Page Description'] = st.text_area(
        "Recommended Page Description",
        help="Enter the meta description. Ensure the primary keyword is targeted and it's 155 characters or less.",
        max_chars=155
    )
    
    st.session_state.data['Recommended URL'] = st.text_input(
        "Recommended URL",
        help="Enter the recommended URL. Ensure the primary keyword is in the slug, it's lowercase, stop-words are omitted, uses hyphens, and is concise."
    )
    
    st.session_state.data['Canonical URL'] = st.session_state.data['Recommended URL']
    
    st.session_state.data['Keyword Included In First 100 Words Of Article'] = st.checkbox(
        "Keyword Included In First 100 Words Of Article",
        help="Check if the primary keyword appears in the first 100 words of the content."
    )
    
    st.session_state.data['Proper Heading Hierarchy'] = st.checkbox(
        "Proper Heading Hierarchy",
        help="Check if the page follows proper heading hierarchy (e.g., single H1, H2s precede H3s, etc.)"
    )

    schema_options = [
        "None", "WebPage", "Article", "Product", "FAQ", "HowTo", "LocalBusiness", 
        "Event", "Recipe", "Review", "BreadcrumbList", "FinancialProduct", 
        "Organization", "Person", "VideoObject", "ImageObject", "Service", 
        "SoftwareApplication", "Course", "JobPosting"
    ]
    
    st.session_state.data['Schema Markup'] = st.multiselect(
        "Schema Markup",
        options=schema_options,
        help="Select the schema markup types that should be included on the page."
    )

    st.session_state.data['Indexable and should be included in sitemap.xml'] = st.checkbox(
        "Indexable and should be included in sitemap.xml",
        help="Check if this page should be indexed and included in the sitemap."
    )

def internal_links():
    st.header("Internal Links")
    
    if 'internal_links' not in st.session_state:
        st.session_state.internal_links = []
    
    for i, link in enumerate(st.session_state.internal_links):
        col1, col2, col3, col4 = st.columns([3, 3, 1, 1])
        with col1:
            link['url'] = st.text_input("URL", link['url'], key=f"url_{i}")
        with col2:
            link['anchor_text'] = st.text_input("Anchor Text", link['anchor_text'], key=f"anchor_{i}")
        with col3:
            link['type'] = st.selectbox("Type", ["Product", "Article"], key=f"type_{i}")
        with col4:
            if st.button("Remove", key=f"remove_link_{i}"):
                st.session_state.internal_links.pop(i)
                st.experimental_rerun()
    
    if st.button("Add Internal Link"):
        st.session_state.internal_links.append({"url": "", "anchor_text": "", "type": "Product"})
        st.experimental_rerun()

def calculate_all_traffic_and_conversions():
    all_traffic_data = []
    if 'keyword_data' in st.session_state and st.session_state.keyword_data:
        for keyword in st.session_state.keyword_data:
            if 'app_start_rate' in st.session_state.data and 'app_submit_rate' in st.session_state.data:
                traffic_data = calculate_traffic_and_conversions(
                    keyword.get('search_volume', 0),
                    keyword.get('current_rank', 1),
                    keyword.get('target_rank', 1),
                    st.session_state.data['app_start_rate'],
                    st.session_state.data['app_submit_rate']
                )
                all_traffic_data.append({
                    'keyword': keyword['keyword'],
                    'incremental_traffic': traffic_data['incremental_traffic'],
                    'incremental_app_starts': traffic_data['incremental_app_starts'],
                    'incremental_app_submits': traffic_data['incremental_app_submits']
                })
    return all_traffic_data

def traffic_and_conversions():
    st.header("Traffic and Conversion Estimation")
    
    if st.session_state.keyword_data:
        for i, keyword in enumerate(st.session_state.keyword_data):
            keyword['target_rank'] = st.number_input(
                f"Target Rank for '{keyword['keyword']}'",
                min_value=1,
                max_value=100,
                value=max(1, keyword['current_rank'] - 1),
                key=f"target_rank_{i}"
            )
        
        st.session_state.data['app_start_rate'] = st.slider(
            "App Start Rate (%)",
            min_value=0.0,
            max_value=100.0,
            value=10.0,
            step=0.1,
            help="Enter the percentage of visitors who start the app."
        ) / 100

        st.session_state.data['app_submit_rate'] = st.slider(
            "App Submit Rate (%)",
            min_value=0.0,
            max_value=100.0,
            value=50.0,
            step=0.1,
            help="Enter the percentage of app starts that result in a submission."
        ) / 100
    else:
        st.write("Please add at least one keyword in the Keyword Research section to enable traffic and conversion estimation.")

def export_to_word(data, keyword_data, traffic_data):
    doc = Document()
    doc.add_heading("SEO Content Optimization Report", 0)
    
    # Keyword Data table
    doc.add_heading("Keyword Data", level=1)
    keyword_table = doc.add_table(rows=1, cols=7)
    keyword_table.style = 'Table Grid'
    hdr_cells = keyword_table.rows[0].cells
    hdr_cells[0].text = 'Keyword'
    hdr_cells[1].text = 'Search Volume'
    hdr_cells[2].text = 'Current Rank'
    hdr_cells[3].text = 'Target Rank'
    hdr_cells[4].text = 'Incremental Traffic'
    hdr_cells[5].text = 'Incremental App Starts'
    hdr_cells[6].text = 'Incremental App Submits'
    
    for kw, td in zip(keyword_data, traffic_data):
        row_cells = keyword_table.add_row().cells
        row_cells[0].text = kw['keyword']
        row_cells[1].text = str(kw['search_volume'])
        row_cells[2].text = str(kw['current_rank'])
        row_cells[3].text = str(kw['target_rank'])
        row_cells[4].text = str(td['incremental_traffic'])
        row_cells[5].text = str(td['incremental_app_starts'])
        row_cells[6].text = str(td['incremental_app_submits'])
    
    # SERP Analysis table
    doc.add_heading("SERP Analysis", level=1)
    serp_table = doc.add_table(rows=1, cols=2)
    serp_table.style = 'Table Grid'
    for field in ['Search Intent For Target Keyword', 'Search Result Features For Target Keyword', 'FAQs Present Within Search Results']:
        row_cells = serp_table.add_row().cells
        row_cells[0].text = field
        row_cells[1].text = data.get(field, '')
    
    # On-Page Elements table
    doc.add_heading("On-Page Elements", level=1)
    elements_table = doc.add_table(rows=1, cols=2)
    elements_table.style = 'Table Grid'
    
    on_page_fields = [
        'Recommended Title (H1)', 'Recommended Page Title', 'Recommended Page Description',
        'Recommended URL', 'Canonical URL', 'Keyword Included In First 100 Words Of Article',
        'Proper Heading Hierarchy', 'Schema Markup', 'Indexable and should be included in sitemap.xml'
    ]
    
    for field in on_page_fields:
        row_cells = elements_table.add_row().cells
        row_cells[0].text = field
        value = data.get(field, '')
        if isinstance(value, bool):
            row_cells[1].text = 'Yes' if value else 'No'
        elif isinstance(value, list):
            row_cells[1].text = ', '.join(value)
        else:
            row_cells[1].text = str(value)
    
    # Internal Links table
    doc.add_heading("Internal Links", level=1)
    links_table = doc.add_table(rows=1, cols=3)
    links_table.style = 'Table Grid'
    hdr_cells = links_table.rows[0].cells
    hdr_cells[0].text = 'URL'
    hdr_cells[1].text = 'Anchor Text'
    hdr_cells[2].text = 'Type'
    
    for link in st.session_state.internal_links:
        row_cells = links_table.add_row().cells
        row_cells[0].text = link['url']
        row_cells[1].text = link['anchor_text']
        row_cells[2].text = link['type']
    
    # Save the document to a BytesIO object
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return doc_io

def get_word_download():
    word_file = export_to_word(st.session_state.data, st.session_state.keyword_data, calculate_all_traffic_and_conversions())
    return word_file.getvalue()

if __name__ == "__main__":
    main()
